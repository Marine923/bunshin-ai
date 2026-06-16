"""Import local files (Markdown, text) into Bunshin storage."""
import sqlite3
from pathlib import Path
from typing import Iterable, Optional

from bunshin.storage import insert_record


DEFAULT_EXTENSIONS = {".md", ".markdown", ".txt", ".pdf", ".docx"}

SKIP_DIRS = {
    "node_modules", ".git", ".venv", "venv", "env", "__pycache__",
    ".pytest_cache", ".ruff_cache", "dist", "build", ".next",
    "target", "Pods", ".cache", "site-packages",
    ".bunshin",  # our own data dir
}

# Patterns matched on file name. Excludes Python packaging artefacts and
# other build outputs that show up as text but aren't real user content.
SKIP_FILE_PATTERNS = {
    "top_level.txt", "requires.txt", "entry_points.txt",
    "dependency_links.txt", "RECORD", "METADATA", "PKG-INFO",
    "package-lock.json", "yarn.lock", "Cargo.lock",
}


def _should_skip_file(path: Path) -> bool:
    if path.name in SKIP_FILE_PATTERNS:
        return True
    # Anything inside a .egg-info, .dist-info, or similar bundle
    if any(part.endswith((".egg-info", ".dist-info", ".egg")) for part in path.parts):
        return True
    return False

MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 MB
CHUNK_SIZE = 1500
CHUNK_OVERLAP = 100


def find_files(
    root: Path,
    extensions: set[str] = DEFAULT_EXTENSIONS,
) -> Iterable[Path]:
    """Walk directory tree, yielding files with matching extensions."""
    if root.is_file():
        if root.suffix.lower() in extensions:
            yield root
        return
    if not root.is_dir():
        return
    try:
        for entry in root.iterdir():
            if entry.name.startswith(".") or entry.name in SKIP_DIRS:
                continue
            if entry.is_dir():
                yield from find_files(entry, extensions)
            elif entry.is_file() and entry.suffix.lower() in extensions:
                if _should_skip_file(entry):
                    continue
                try:
                    if entry.stat().st_size <= MAX_FILE_SIZE:
                        yield entry
                except OSError:
                    continue
    except (OSError, PermissionError):
        return


def read_text(path: Path) -> Optional[str]:
    """Read file text. Dispatches on extension for PDF / DOCX, else
    decodes as UTF-8 (with Shift-JIS fallback for legacy Japanese files)."""
    ext = path.suffix.lower()
    if ext == ".pdf":
        return _read_pdf(path)
    if ext == ".docx":
        return _read_docx(path)
    # Plain text fallback (md, txt, anything else with text)
    for encoding in ("utf-8", "cp932"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
        except OSError:
            return None
    return None


def _read_pdf(path: Path) -> Optional[str]:
    """Extract text from a PDF. Returns None on unreadable / encrypted."""
    try:
        from pypdf import PdfReader
    except ImportError:
        return None
    try:
        reader = PdfReader(str(path))
        if reader.is_encrypted:
            try:
                reader.decrypt("")  # try empty password
            except Exception:
                return None
        parts = []
        for i, page in enumerate(reader.pages):
            try:
                txt = page.extract_text() or ""
            except Exception:
                continue
            if txt.strip():
                parts.append(f"[p.{i+1}]\n{txt.strip()}")
        return "\n\n".join(parts) if parts else None
    except Exception:
        return None


def _read_docx(path: Path) -> Optional[str]:
    """Extract text from a .docx file."""
    try:
        import docx  # python-docx
    except ImportError:
        return None
    try:
        doc = docx.Document(str(path))
        parts = []
        for p in doc.paragraphs:
            t = p.text.strip()
            if t:
                parts.append(t)
        # Tables: serialize cells as tab-separated rows
        for table in doc.tables:
            for row in table.rows:
                row_text = "\t".join(cell.text.strip() for cell in row.cells)
                if row_text.strip("\t"):
                    parts.append(row_text)
        return "\n".join(parts) if parts else None
    except Exception:
        return None


def chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[str]:
    """Split text into chunks, preferring paragraph boundaries."""
    text = text.strip()
    if len(text) <= chunk_size:
        return [text] if text else []

    chunks = []
    start = 0
    while start < len(text):
        end = min(start + chunk_size, len(text))
        if end < len(text):
            # Try to break at paragraph (\n\n) within last ~200 chars
            window_start = max(end - 200, start + 100)
            paragraph = text.rfind("\n\n", window_start, end)
            if paragraph != -1:
                end = paragraph + 2
            else:
                # Fall back to single newline
                newline = text.rfind("\n", window_start, end)
                if newline != -1:
                    end = newline + 1
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= len(text):
            break
        start = max(end - overlap, start + 1)
    return chunks


def delete_file_records(conn: sqlite3.Connection, source_id: str) -> None:
    cursor = conn.execute(
        "SELECT id FROM records WHERE source = 'file' AND source_id = ?",
        (source_id,),
    )
    ids = [row[0] for row in cursor.fetchall()]
    if not ids:
        return
    placeholders = ",".join(["?"] * len(ids))
    try:
        conn.execute(
            f"DELETE FROM records_vec WHERE record_id IN ({placeholders})",
            ids,
        )
    except sqlite3.OperationalError:
        pass
    conn.execute(
        f"DELETE FROM records WHERE id IN ({placeholders})",
        ids,
    )
    conn.commit()


def get_last_mtime(conn: sqlite3.Connection, path_str: str) -> Optional[int]:
    cursor = conn.execute(
        "SELECT value FROM settings WHERE key = ?",
        (f"file_mtime:{path_str}",),
    )
    row = cursor.fetchone()
    return int(row[0]) if row else None


def set_last_mtime(conn: sqlite3.Connection, path_str: str, mtime: int) -> None:
    conn.execute(
        "INSERT OR REPLACE INTO settings(key, value) VALUES(?, ?)",
        (f"file_mtime:{path_str}", str(mtime)),
    )


def import_files(
    conn: sqlite3.Connection,
    root: Path,
    extensions: Optional[set[str]] = None,
    force: bool = False,
    verbose: bool = False,
) -> dict:
    """Import all matching files under root. Skips files unchanged since last import."""
    try:
        from bunshin.storage import load_vec_extension
        load_vec_extension(conn)
    except Exception:
        pass

    extensions = extensions or DEFAULT_EXTENSIONS
    stats = {
        "files_scanned": 0,
        "files_unchanged": 0,
        "files_reimported": 0,
        "files_failed": 0,
        "chunks_inserted": 0,
    }

    for path in find_files(root, extensions):
        stats["files_scanned"] += 1
        path_str = str(path)
        try:
            current_mtime = int(path.stat().st_mtime)
        except OSError:
            stats["files_failed"] += 1
            continue

        last_mtime = get_last_mtime(conn, path_str)
        if not force and last_mtime is not None and last_mtime >= current_mtime:
            stats["files_unchanged"] += 1
            continue

        text = read_text(path)
        if not text:
            stats["files_failed"] += 1
            continue

        # Re-import: wipe old records for this file
        delete_file_records(conn, path_str)

        chunks = chunk_text(text)
        for i, chunk in enumerate(chunks):
            insert_record(
                conn,
                source="file",
                timestamp=current_mtime,
                content=chunk,
                source_id=path_str,
                metadata={
                    "path": path_str,
                    "name": path.name,
                    "ext": path.suffix.lower(),
                    "chunk_index": i,
                    "chunk_count": len(chunks),
                },
                file_path=path_str,
            )
            stats["chunks_inserted"] += 1

        set_last_mtime(conn, path_str, current_mtime)
        conn.commit()
        stats["files_reimported"] += 1
        if verbose:
            print(f"Imported: {path} ({len(chunks)} chunks)")

    return stats
